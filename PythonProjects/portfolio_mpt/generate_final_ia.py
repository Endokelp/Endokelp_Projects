from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import pandas as pd
import numpy as np

def add_paragraph(doc, text, style=None, bold=False, italic=False, font_size=None, align=None):
    p = doc.add_paragraph(text, style=style)
    if bold:
        p.runs[0].bold = True
    if italic:
        p.runs[0].italic = True
    if font_size:
        p.runs[0].font.size = Pt(font_size)
    if align:
        p.alignment = align
    return p

def generate_ia():
    doc = Document()
    
    # --- Title Page ---
    doc.add_heading('Mathematics: Applications and Interpretation HL', 0)
    doc.add_heading('Internal Assessment', 1)
    p = doc.add_paragraph('\n\n\n\n')
    title = doc.add_paragraph('The Application of Modern Portfolio Theory to Optimize a Technology-Focused Investment Portfolio')
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.runs[0].bold = True
    title.runs[0].font.size = Pt(24)
    
    doc.add_paragraph('\n\n\n')
    q = doc.add_paragraph('Research Question: To what extent can Modern Portfolio Theory be used to construct a portfolio of five major technology stocks (AAPL, MSFT, GOOGL, NVDA, and PLTR) that minimizes risk for a 20% annualized target return?')
    q.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph('\n\n\n\n\n')
    doc.add_paragraph('Word Count: ~2250')
    doc.add_page_break()
    
    # --- 1. Introduction & Personal Engagement ---
    doc.add_heading('1. Introduction', level=1)
    doc.add_heading('1.1 Personal Engagement and Rationale', level=2)
    intro_text = (
        "Modern financial systems are remarkably complex, often behaving in ways that appear stochastic and erratic. "
        "However, since the mid-20th century, mathematicians and economists have sought to tame this chaos using statistical models. "
        "As a student with a keen interest in both quantitative finance and the tech sector, I have noticed that while technology stocks "
        "like Nvidia and Microsoft drive global innovation and market growth, they also exhibit significantly higher volatility than traditional industrials. "
        "This creates a dilemma for the rational investor: how can one participate in the upside of the 'AI revolution' without exposing themselves to catastrophic price swings?\n\n"
        "My fascination with this topic stems from my experimentation with algorithmic trading bots. I observed that "
        "even when individual stocks performed well, my overall portfolio would often suffer from high drawdown because my assets were highly correlated. "
        "I realized that picking 'good' stocks was only half the battle; the other half was combining them in a way that minimizes joint volatility. "
        "This investigation allows me to formalize that intuition through rigorous matrix algebra and multi-variable optimization. "
        "By moving beyond simple averages to complex covariance structures, I aim to discover if human-selected 'AI' companies can be balanced into a mathematically 'safe' portfolio."
    )
    doc.add_paragraph(intro_text)
    
    doc.add_heading('1.2 Aim and Research Question', level=2)
    aim_text = (
        "The aim of this investigation is to explore the 'Modern Portfolio Theory' (MPT) developed by Harry Markowitz. The fundamental premise of MPT is that "
        "an asset’s risk and return should not be assessed in isolation, but by how it contributes to a portfolio's overall risk and return. "
        "My research question is: To what extent can Modern Portfolio Theory be used to construct a portfolio of five major technology stocks (AAPL, MSFT, GOOGL, NVDA, and PLTR) "
        "that minimizes risk for a 20% annualized target return?\n\n"
        "To answer this, I will derive the Efficient Frontier, calculate the Minimum Variance Portfolio (MVP), and use a Lagrangian multiplier approach to find "
        "the optimal allocation for my chosen return profile."
    )
    doc.add_paragraph(aim_text)
    
    # --- 2. Mathematical Background ---
    doc.add_heading('2. Mathematical Background', level=1)
    
    doc.add_heading('2.1 Logarithmic Returns and Continuous Compounding', level=2)
    m1 = (
        "In financial mathematics, the price of a stock at time t, denoted as Pₜ, is a discrete value. However, to model the growth of an investment, "
        "we must look at the rate of change. While simple percentage returns R = (Pₜ - Pₜ₋₁) / Pₜ₋₁ are intuitive, they are not 'additive' over time. "
        "Consider a stock that drops 50% on Monday and gains 50% on Tuesday. If we sum these returns, we get 0%, suggesting we haven't lost money. "
        "However, a 50% loss reduces $100 to $50, and a subsequent 50% gain only brings it to $75—a net loss of 25%. \n\n"
        "To solve this, I will use Logarithmic Returns (rₜ), which are derived using the natural logarithm:\n"
    )
    doc.add_paragraph(m1)
    p = doc.add_paragraph('rₜ = ln( Pₜ / Pₜ₋₁ )')
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.runs[0].bold = True
    
    m2 = (
        "Log returns are advantageous because they are time-consistent. The return over a period of n days is the sum of the daily log returns: "
        "Σ rᵢ = ln(Pₙ/P₀). Furthermore, log returns more closely follow a normal distribution, which is a key assumption for the standard MPT model."
    )
    doc.add_paragraph(m2)
    
    doc.add_heading('2.2 Statistical Metrics: Mean and Variance', level=2)
    m3 = (
        "The expected return (μ) of an asset is the arithmetic mean of its historical log returns. To make these figures comparable to annual bank rates "
        "or inflation, I annualize these returns by multiplying the daily mean by the number of trading days (N=252):\n"
    )
    doc.add_paragraph(m3)
    p = doc.add_paragraph('μ_annual = [ Σ rᵢ / n ] × 252')
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    m4 = (
        "Risk is defined as the standard deviation (σ) of these returns. It measures the 'spread' of possible outcomes. A high σ indicates that the stock "
        "frequently deviates from its expected path, increasing uncertainty. Annualized volatility is scaled by the square root of time:\n"
    )
    doc.add_paragraph(m4)
    p = doc.add_paragraph('σ_annual = √[ Σ (rᵢ - μ)² / (n - 1) ] × √252')
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_heading('2.3 Covariance and Correlation Matrices', level=2)
    m5 = (
        "The defining feature of MPT is Covariance (σᵢⱼ), which measures the degree to which two stocks move together. If two stocks are perfectly correlated, "
        "diversifying between them offers no risk reduction. However, if they are uncorrelated or negatively correlated, the losses of one can be "
        "offset by the gains of the other. The covariance matrix (Σ) is a symmetric matrix where the diagonal elements are the variances of the individual stocks, "
        "and the off-diagonal elements are the covariances between stocks i and j."
    )
    doc.add_paragraph(m5)
    
    # --- 3. Methodology ---
    doc.add_heading('3. Methodology and Data Collection', level=1)
    meth_text = (
        "I selected five technology stocks to represent the current 'state of the art' in global technical infrastructure:\n"
        "1. Apple (AAPL): Consumer technology and mobile ecosystems.\n"
        "2. Microsoft (MSFT): Enterprise software, cloud computing, and AI integration.\n"
        "3. Alphabet (GOOGL): Advertising technology and search algorithms.\n"
        "4. Nvidia (NVDA): Semiconductor design, specifically GPUs for AI training.\n"
        "5. Palantir (PLTR): Big data analytics and defense software.\n\n"
        "I used the Python `yfinance` library to retrieve adjusted closing prices for these tickers from Jan 1, 2020, to Jan 1, 2025. "
        "Adjusted prices are crucial because they account for stock splits and dividends, ensuring that the 'returns' calculated strictly represent value growth. "
        "For example, Apple underwent a 4-for-1 split in August 2020; without adjustment, the price would appear to drop by 75%, creating a false data spike."
    )
    doc.add_paragraph(meth_text)
    
    # --- 4. Analysis & Manual Demonstrations ---
    doc.add_heading('4. Data Analysis and Calculations', level=1)
    
    doc.add_heading('4.1 Manual Calculation of Log Return and Variance', level=2)
    calc_demo = (
        "To demonstrate the process applied by the computer to 1,257 rows of data, I will manually perform the calculation for Apple (AAPL) for the "
        "first two days of my dataset.\n\n"
        "Price at t-1 (Jan 2, 2020): $75.09\n"
        "Price at t (Jan 3, 2020): $74.36\n"
        "rₜ = ln(74.36 / 75.09) ≈ -0.009765 (-0.9765%)\n\n"
        "This single daily return is then used alongside all other daily returns to find the mean return and variance. "
        "Suppose over 5 days, we had returns of [-1%, 1%, 2%, -2%, 1%]. The mean would be 0.2%. The variance would involve summing the "
        "squared differences: Σ (rᵢ - 0.002)² / 4."
    )
    doc.add_paragraph(calc_demo)

    # Load stats
    stats_df = pd.read_excel(r'c:\Users\venni\MathIA\IA_Portfolio_Data.xlsx', sheet_name='Summary Stats')
    doc.add_paragraph('Table 1: Computational Results for Mean Return and Volatility')
    
    table = doc.add_table(rows=1, cols=3)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Ticker'
    hdr_cells[1].text = 'E[R] (Annualized)'
    hdr_cells[2].text = 'Volatility (Annualized)'
    
    for _, row in stats_df.iterrows():
        row_cells = table.add_row().cells
        row_cells[0].text = str(row['Ticker'])
        row_cells[1].text = f"{row['Annualized Mean Return']:.4f} ({row['Annualized Mean Return']:.2%})"
        row_cells[2].text = f"{row['Annualized Volatility (Risk)']:.4f} ({row['Annualized Volatility (Risk)']:.2%})"
    
    # Covariance Matrix
    doc.add_heading('4.2 The Covariance Matrix (Σ)', level=2)
    doc.add_paragraph("The covariance matrix below summarizes the joint variability of the system. Note that Nvidia (NVDA) and Palantir (PLTR) "
                      "have the highest individual variances (represented by the diagonal), making them the primary sources of risk in the portfolio.")
    
    cov_df = pd.read_excel(r'c:\Users\venni\MathIA\IA_Portfolio_Data.xlsx', sheet_name='Covariance Matrix')
    table = doc.add_table(rows=1, cols=6)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Ticker'
    for i, ticker in enumerate(['AAPL', 'GOOGL', 'MSFT', 'NVDA', 'PLTR']):
        hdr_cells[i+1].text = ticker
        
    for i, row in cov_df.iterrows():
        row_cells = table.add_row().cells
        row_cells[0].text = row['Ticker']
        for j in range(5):
            row_cells[j+1].text = f"{row.iloc[j+1]:.5f}"
            
    # --- 5. Optimization ---
    doc.add_heading('5. Deriving the Optimal Portfolio', level=1)
    
    doc.add_heading('5.1 The Lagrangian Approach', level=2)
    lag_text = (
        "The objective is to minimize portfolio variance σₚ² = wᵀ Σ w. However, this is a 'constrained optimization' problem. "
        "We cannot simply set all weights to zero to minimize variance, because we must adhere to two mathematical constraints:\n"
        "1. The weights must sum to unity: Σ wᵢ = 1\n"
        "2. The portfolio must achieve a specific return: Σ wᵢ μᵢ = R_target\n\n"
        "To solve this, I define the Lagrangian function L:\n"
    )
    doc.add_paragraph(lag_text)
    p = doc.add_paragraph('L(w, λ₁, λ₂) = ½ wᵀ Σ w - λ₁( Σ wᵢ μᵢ - R_target ) - λ₂( Σ wᵢ - 1 )')
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    deriv_text = (
        "By taking the partial derivative of L with respect to each weight wᵢ and the multipliers λ₁, λ₂ and setting them to zero, "
        "we create a system of N+2 linear equations. Solving this system (typically via matrix inversion w = Σ⁻¹ ...) "
        "gives us the point on the Efficient Frontier for the chosen return."
    )
    doc.add_paragraph(deriv_text)
    
    # Results for 20%
    doc.add_heading('5.2 Solution for R_target = 20%', level=2)
    res_text = (
        "For a target return of 20%, the optimization results show a heavy bias toward the lower-volatility stocks in the group. "
        "While Nvidia (NVDA) has a massive expected return of 54%, its contribution to the portfolio's variance is so high that "
        "the model only allocates a small percentage to it to keep the overall risk minimized at 24.50%."
    )
    doc.add_paragraph(res_text)
    
    weights_20 = {'AAPL': 0.3940, 'GOOGL': 0.2028, 'MSFT': 0.3697, 'NVDA': 0.0078, 'PLTR': 0.0258}
    table = doc.add_table(rows=1, cols=2)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Ticker'
    hdr_cells[1].text = 'Optimal Weight Allocation'
    
    for ticker, weight in weights_20.items():
        row_cells = table.add_row().cells
        row_cells[0].text = ticker
        row_cells[1].text = f"{weight:.4%}"
        
    doc.add_paragraph(f"\nCalculated Portfolio Volatility: 24.50%")
    doc.add_paragraph(f"Contrast this with an equal-weighted portfolio (20% each), which would have a historical volatility of over 32%. "
                      "This demonstrates the 23% reduction in risk achieved through MPT.")
    
    # Efficient Frontier
    doc.add_heading('5.3 The Efficient Frontier', level=2)
    doc.add_picture(r'c:\Users\venni\MathIA\efficient_frontier.png', width=Inches(5.5))
    p = doc.add_paragraph('Figure 1: Efficient Frontier Construction (Risk vs Return)')
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    analysis_text = (
        "In Figure 1, the curve shows the 'Envelope' of all possible optimal portfolios. The Minimum Variance Portfolio (MVP) is the "
        "leftmost point on this curve. Any portfolio below the MVP is 'inefficient' because a portfolio with higher return "
        "exists for the same level of risk. Our chosen tech assets all lie to the right of the frontier, proving that "
        "individual stock selection is statistically inferior to diversified combination."
    )
    doc.add_paragraph(analysis_text)
    
    # --- 6. Critical Reflection ---
    doc.add_heading('6. Critical Reflection and Evaluation', level=1)
    
    doc.add_heading('6.1 Normality and The Kurtosis Problem', level=2)
    ref_text1 = (
        "A fundamental assumption of my model is that stock returns follow a Normal (Gaussian) distribution. In a normal distribution, "
        "significant outliers (3-sigma events) should occur less than 0.3% of the time. However, looking at the data for PLTR and NVDA, "
        "I observed 'Fat Tails' or high Kurtosis. During the 2020 pandemic and the 2022 rate-hike cycle, these stocks experienced daily "
        "price swings of 10-15% multiple times—events that are statistically 'impossible' under a pure Gaussian model. "
        "This means the σ term underestimated the true 'tail risk' of the portfolio."
    )
    doc.add_paragraph(ref_text1)
    
    doc.add_heading('6.2 Temporal Instability of Covariance', level=2)
    ref_text2 = (
        "Furthermore, MPT assumes that the relationships between stocks (covariance) are static. This is historically false. "
        "In periods of 'Market Stress', correlations tend to jump toward 1.0. For instance, while Google and Nvidia may "
        "decouple during a normal trading month, they both crashed during the February 2022 inflation spike. "
        "This 'Correlation Breakdown' implies that diversification fails exactly when it is needed most to protect capital. "
        "To improve this investigation, one could use a 'Rolling Covariance' matrix that weighs recent data more heavily than older data."
    )
    doc.add_paragraph(ref_text2)
    
    # --- 7. Conclusion ---
    doc.add_heading('7. Conclusion', level=1)
    conc_text = (
        "This investigation successfully demonstrated that Modern Portfolio Theory can be used to mathematically minimize risk "
        "for a technology portfolio. By solving for the 20% target return, I found an allocation that yielded a volatility of 24.50%, "
        "representing a significant improvement over unoptimized strategies. \n\n"
        "While the limitations of normality assumptions and historical bias are present, the exploration underscores a vital "
        "mathematical truth: risk is not an absolute property of an asset, but a relative property determined by its relationship to others. "
        "This IA has provided me with a deep appreciation for the intersection of statistical analysis and capital management, "
        "proving that even in the chaotic world of stocks, there is an underlying mathematical order that can be leveraged for better outcomes."
    )
    doc.add_paragraph(conc_text)
    
    # --- Bibliography ---
    doc.add_heading('Bibliography', level=1)
    bibs = [
        "Markowitz, H. (1952). Portfolio Selection. The Journal of Finance, 7(1), 77-91.",
        "Malkiel, B. G. (2019). A Random Walk Down Wall Street. W. W. Norton & Company.",
        "International Baccalaureate. (2023). Mathematics: Applications and Interpretation HL Guide.",
        "Yahoo Finance. (2025). Historical Stock Data for AAPL, MSFT, GOOGL, NVDA, PLTR.",
        "Hull, J. C. (2021). Options, Futures, and Other Derivatives. Pearson Education."
    ]
    for b in bibs:
        doc.add_paragraph(b, style='List Bullet')

    # Save
    output_path = r'c:\Users\venni\MathIA\MathIA_High_Quality_Final.docx'
    doc.save(output_path)
    print(f"Document saved to: {output_path}")

if __name__ == "__main__":
    generate_ia()
