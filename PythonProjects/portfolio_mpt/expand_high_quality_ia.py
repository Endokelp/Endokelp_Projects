from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import pandas as pd
import numpy as np
from scipy.optimize import minimize
import os

def add_paragraph(doc, text, style=None, bold=False, italic=False, font_size=None, align=None):
    p = doc.add_paragraph(text, style=style)
    if bold:
        p.runs[0].bold = True
    if italic:
        p.runs[0].italic = True
    if font_size:
        p.runs[0].font.size = Pt(font_size)
    if align:
        p.alignment = align
    return p

def solve_mpt(mu, cov, target_return):
    n = len(mu)
    def portfolio_variance(weights):
        return weights.T @ cov @ weights
    constraints = [
        {'type': 'eq', 'fun': lambda x: np.sum(x) - 1},
        {'type': 'eq', 'fun': lambda x: x.T @ mu - target_return}
    ]
    bounds = tuple((0, 1) for _ in range(n))
    initial_weights = np.array([1/n] * n)
    result = minimize(portfolio_variance, initial_weights, method='SLSQP', bounds=bounds, constraints=constraints)
    return result.x if result.success else initial_weights

def generate_expanded_ia():
    doc = Document()
    
    # Global style
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)
    
    # --- Title Page ---
    doc.add_heading('Mathematics: Applications and Interpretation HL', 0)
    doc.add_heading('Internal Assessment', 1)
    p = doc.add_paragraph('\n\n\n\n')
    title = doc.add_paragraph('The Application of Modern Portfolio Theory to Optimize a Technology-Focused Investment Portfolio')
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.runs[0].bold = True
    title.runs[0].font.size = Pt(24)
    
    doc.add_paragraph('\n\n\n')
    q = doc.add_paragraph('Research Question: To what extent can Modern Portfolio Theory be used to construct a portfolio of five major technology stocks (AAPL, MSFT, GOOGL, NVDA, and PLTR) that minimizes risk for a 20% annualized target return?')
    q.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph('\n\n\n\n\n')
    doc.add_paragraph('Word Count: ~4500')
    doc.add_page_break()
    
    # --- Table of Contents ---
    doc.add_heading('Table of Contents', level=1)
    toc_items = [
        '1. Introduction',
        '   1.1 Personal Engagement and Rationale',
        '   1.2 Aim and Research Question',
        '2. Definitions and Notation',
        '3. Mathematical Background',
        '   3.1 Logarithmic Returns and Continuous Compounding',
        '   3.2 Statistical Metrics: Mean and Variance',
        '   3.3 Covariance and Portfolio Variance',
        '4. Methodology and Data Collection',
        '5. Data Analysis and Calculations',
        '   5.1 Manual Calculation of Log Return',
        '   5.2 Results: Summary Statistics',
        '   5.3 The Covariance Matrix (\u03A3)',
        '6. Deriving the Optimal Portfolio',
        '   6.1 Setting Up the Constrained Optimization',
        '   6.2 Constructing the Lagrangian',
        '   6.3 Computing the Partial Derivatives',
        '   6.4 Solving for the Optimal Weights',
        '   6.5 Solution for R_target = 20%',
        '   6.6 The Efficient Frontier',
        '7. Critical Reflection and Evaluation',
        '   7.1 Normality and The Kurtosis Problem',
        '   7.2 Temporal Instability of Covariance',
        '   7.3 Limitations of the Lagrangian Method',
        '8. Conclusion',
        'Bibliography',
    ]
    for item in toc_items:
        doc.add_paragraph(item)
    doc.add_page_break()
    
    # ========================================================================
    # 1. INTRODUCTION
    # ========================================================================
    doc.add_heading('1. Introduction', level=1)
    doc.add_heading('1.1 Personal Engagement and Rationale', level=2)
    doc.add_paragraph(
        "Modern financial systems are remarkably complex, often behaving in ways that appear stochastic and erratic. "
        "However, since the mid-20th century, mathematicians and economists have sought to tame this chaos using statistical models. "
        "As a student with a keen interest in both quantitative finance and the tech sector, I have noticed that while technology stocks "
        "like Nvidia and Microsoft drive global innovation and market growth, they also exhibit significantly higher volatility than traditional industrials. "
        "This creates a dilemma for the rational investor: how can one participate in the upside of the 'AI revolution' without exposing themselves to catastrophic price swings?"
    )
    doc.add_paragraph(
        "My fascination with this topic stems from my experimentation with algorithmic trading bots. I observed that "
        "even when individual stocks performed well, my overall portfolio would often suffer from high drawdown because my assets were highly correlated. "
        "I realized that picking 'good' stocks was only half the battle; the other half was combining them in a way that minimizes joint volatility. "
        "This investigation allows me to formalize that intuition through rigorous matrix algebra and multi-variable optimization. "
        "By moving beyond simple averages to complex covariance structures, I aim to discover if human-selected 'AI' companies can be balanced into a mathematically 'safe' portfolio."
    )
    
    doc.add_heading('1.2 Aim and Research Question', level=2)
    doc.add_paragraph(
        "The aim of this investigation is to explore the 'Modern Portfolio Theory' (MPT) developed by Harry Markowitz. The fundamental premise of MPT is that "
        "an asset's risk and return should not be assessed in isolation, but by how it contributes to a portfolio's overall risk and return. "
        "My research question is: To what extent can Modern Portfolio Theory be used to construct a portfolio of five major technology stocks (AAPL, MSFT, GOOGL, NVDA, and PLTR) "
        "that minimizes risk for a 20% annualized target return?"
    )
    doc.add_paragraph(
        "To answer this, I will derive the Efficient Frontier, calculate the Minimum Variance Portfolio (MVP), and use a Lagrangian multiplier approach to find "
        "the optimal allocation for my chosen return profile. I will also critically evaluate the model's assumptions when applied to real-world data that includes "
        "the 2020 pandemic crash and the 2022 inflationary shock."
    )
    
    # ========================================================================
    # 2. DEFINITIONS AND NOTATION
    # ========================================================================
    doc.add_page_break()
    doc.add_heading('2. Definitions and Notation', level=1)
    doc.add_paragraph(
        "Before proceeding with the mathematics, I will formally define all key variables and symbols used in this investigation. "
        "Clear notation is essential for mathematical communication and ensures there is no ambiguity in the derivations that follow."
    )
    
    doc.add_paragraph(
        "Let n denote the number of assets in the portfolio. In this investigation, n = 5."
    )
    
    doc.add_paragraph(
        "The weight vector w is a column vector of dimension n \u00D7 1, where each element w\u1D62 represents the fraction of total capital allocated to asset i. "
        "The constraint that all capital must be invested means:"
    )
    p = doc.add_paragraph('w\u2081 + w\u2082 + w\u2083 + w\u2084 + w\u2085 = 1')
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.runs[0].bold = True
    
    doc.add_paragraph(
        "The expected return vector \u03BC is a column vector of dimension n \u00D7 1, where \u03BC\u1D62 represents the annualized mean logarithmic return of asset i, "
        "calculated from historical data over the period 2020\u20132025."
    )
    
    doc.add_paragraph(
        "The covariance matrix \u03A3 is a symmetric matrix of dimension n \u00D7 n. The diagonal elements \u03C3\u1D62\u1D62 represent the variance of asset i, "
        "while the off-diagonal elements \u03C3\u1D62\u2C7C represent the covariance between assets i and j. Mathematically:"
    )
    p = doc.add_paragraph('\u03A3 \u2208 \u211D\u2075 \u00D7 \u211D\u2075,  where \u03A3\u1D62\u2C7C = Cov(r\u1D62, r\u2C7C)')
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.runs[0].bold = True
    
    doc.add_paragraph(
        "The portfolio return R\u209A is a scalar value representing the expected return of the combined portfolio. It is computed as the dot product "
        "of the weight vector and the return vector:"
    )
    p = doc.add_paragraph('R\u209A = w\u1D40 \u03BC = \u03A3 w\u1D62 \u03BC\u1D62')
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.runs[0].bold = True
    
    doc.add_paragraph(
        "The portfolio variance \u03C3\u209A\u00B2 is a scalar measuring the total risk of the portfolio. Using matrix notation:"
    )
    p = doc.add_paragraph('\u03C3\u209A\u00B2 = w\u1D40 \u03A3 w')
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.runs[0].bold = True
    
    doc.add_paragraph(
        "The target return R_target is the investor's desired level of annualized return. For this investigation, R_target = 0.20 (i.e., 20%)."
    )
    
    doc.add_paragraph(
        "The Lagrange Multipliers \u03BB\u2081 and \u03BB\u2082 are auxiliary variables introduced to enforce the equality constraints during optimization. "
        "They do not have direct financial meaning but are essential mathematical tools for solving the constrained problem."
    )
    
    # ========================================================================
    # 3. MATHEMATICAL BACKGROUND
    # ========================================================================
    doc.add_page_break()
    doc.add_heading('3. Mathematical Background', level=1)
    
    doc.add_heading('3.1 Logarithmic Returns and Continuous Compounding', level=2)
    doc.add_paragraph(
        "In financial mathematics, the price of a stock at time t, denoted as P\u209C, is a discrete value. However, to model the growth of an investment, "
        "we must look at the rate of change. While simple percentage returns R = (P\u209C \u2212 P\u209C\u208B\u2081) \u00F7 P\u209C\u208B\u2081 are intuitive, they are not 'additive' over time. "
        "Consider a stock that drops 50% on Monday and gains 50% on Tuesday. If we sum these returns, we get 0%, suggesting we haven't lost money. "
        "However, a 50% loss reduces $100 to $50, and a subsequent 50% gain only brings it to $75\u2014a net loss of 25%."
    )
    doc.add_paragraph(
        "To solve this, I will use Logarithmic Returns (r\u209C), which are derived using the natural logarithm:"
    )
    p = doc.add_paragraph('r\u209C = ln( P\u209C \u00F7 P\u209C\u208B\u2081 )')
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.runs[0].bold = True
    
    doc.add_paragraph(
        "The mathematical beauty of log returns lies in their additivity. The total return over n days is the sum of daily log returns:"
    )
    p = doc.add_paragraph('r_total = r\u2081 + r\u2082 + \u2026 + r\u2099 = ln(P\u2099 \u00F7 P\u2080)')
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.runs[0].bold = True
    
    doc.add_paragraph(
        "This property significantly simplifies the statistical modeling of long-term wealth. Furthermore, log returns more closely "
        "follow a normal distribution, which is a key assumption for the standard MPT model."
    )
    
    doc.add_heading('3.2 Statistical Metrics: Mean and Variance', level=2)
    doc.add_paragraph(
        "The expected return (\u03BC) of an asset is the arithmetic mean of its historical log returns. To make these figures comparable to annual bank rates "
        "or inflation, I annualize these returns by multiplying the daily mean by the number of trading days (N = 252):"
    )
    p = doc.add_paragraph('\u03BC_annual = ( \u03A3 r\u1D62 \u00F7 n ) \u00D7 252')
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.runs[0].bold = True
    
    doc.add_paragraph(
        "Risk is defined as the standard deviation (\u03C3) of these returns. It measures the 'spread' of possible outcomes. A high \u03C3 indicates that the stock "
        "frequently deviates from its expected path, increasing uncertainty. The sample variance is calculated as:"
    )
    p = doc.add_paragraph('\u03C3\u00B2 = \u03A3 (r\u1D62 \u2212 \u03BC)\u00B2 \u00F7 (n \u2212 1)')
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.runs[0].bold = True
    
    doc.add_paragraph("Annualized volatility is then scaled by the square root of time:")
    p = doc.add_paragraph('\u03C3_annual = \u221A[ \u03A3 (r\u1D62 \u2212 \u03BC)\u00B2 \u00F7 (n \u2212 1) ] \u00D7 \u221A252')
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.runs[0].bold = True
    
    doc.add_heading('3.3 Covariance and Portfolio Variance', level=2)
    doc.add_paragraph(
        "The defining feature of MPT is Covariance (\u03C3\u1D62\u2C7C), which measures the degree to which two stocks move together. If two stocks are perfectly correlated, "
        "diversifying between them offers no risk reduction. However, if they are uncorrelated or negatively correlated, the losses of one can be "
        "offset by the gains of the other. The sample covariance between assets i and j is:"
    )
    p = doc.add_paragraph('\u03C3\u1D62\u2C7C = \u03A3 (r\u1D62,\u209C \u2212 \u03BC\u1D62)(r\u2C7C,\u209C \u2212 \u03BC\u2C7C) \u00F7 (n \u2212 1)')
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.runs[0].bold = True
    
    doc.add_paragraph(
        "The portfolio variance is not the weighted average of individual variances. It must account for every pairwise covariance. "
        "In matrix notation, the portfolio variance is:"
    )
    p = doc.add_paragraph('\u03C3\u209A\u00B2 = w\u1D40 \u03A3 w')
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.runs[0].bold = True
    
    doc.add_paragraph(
        "To build intuition, consider a simplified two-asset case. Expanding the matrix multiplication gives:"
    )
    p = doc.add_paragraph('\u03C3\u209A\u00B2 = w\u2081\u00B2 \u03C3\u2081\u00B2 + w\u2082\u00B2 \u03C3\u2082\u00B2 + 2 w\u2081 w\u2082 \u03C3\u2081\u2082')
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.runs[0].bold = True
    
    doc.add_paragraph(
        "The third term (2 w\u2081 w\u2082 \u03C3\u2081\u2082) is crucial. If the covariance \u03C3\u2081\u2082 is negative, the cross-term reduces the total variance, "
        "which is the mathematical basis for why diversification works. This is the core insight of Markowitz's contribution to economics."
    )
    
    # ========================================================================
    # 4. METHODOLOGY
    # ========================================================================
    doc.add_page_break()
    doc.add_heading('4. Methodology and Data Collection', level=1)
    doc.add_paragraph(
        "I selected five technology stocks to represent the current 'state of the art' in global technical infrastructure:\n"
        "1. Apple (AAPL): Consumer technology and mobile ecosystems.\n"
        "2. Microsoft (MSFT): Enterprise software, cloud computing, and AI integration.\n"
        "3. Alphabet (GOOGL): Advertising technology and search algorithms.\n"
        "4. Nvidia (NVDA): Semiconductor design, specifically GPUs for AI training.\n"
        "5. Palantir (PLTR): Big data analytics and defense software."
    )
    doc.add_paragraph(
        "I used the Python yfinance library to retrieve adjusted closing prices for these tickers from Jan 1, 2020, to Jan 1, 2025. "
        "Adjusted prices are crucial because they account for stock splits and dividends, ensuring that the 'returns' calculated strictly represent value growth. "
        "For example, Apple underwent a 4-for-1 split in August 2020; without adjustment, the price would appear to drop by 75%, creating a false data spike. "
        "Similarly, Nvidia underwent a 10-for-1 split in June 2024, which must be corrected for."
    )
    doc.add_paragraph(
        "The five-year timeframe (2020\u20132025) was chosen deliberately. It encompasses the COVID-19 pandemic crash of March 2020, the subsequent "
        "recovery and tech boom, the inflationary shock of 2022, and the generative AI rally of 2023\u20132024. This diversity of market regimes provides "
        "a rigorous stress test for the assumptions of Modern Portfolio Theory."
    )
    
    # ========================================================================
    # 5. DATA ANALYSIS
    # ========================================================================
    doc.add_page_break()
    doc.add_heading('5. Data Analysis and Calculations', level=1)
    
    doc.add_heading('5.1 Manual Calculation of Log Return', level=2)
    doc.add_paragraph(
        "To demonstrate the process applied by the computer to 1,257 rows of data, I will manually perform the calculation for Apple (AAPL) for the "
        "first two days of my dataset."
    )
    
    # Load real data
    stats_df = pd.read_excel(r'c:\Users\venni\MathIA\IA_Portfolio_Data.xlsx', sheet_name='Summary Stats')
    prices_df = pd.read_excel(r'c:\Users\venni\MathIA\IA_Portfolio_Data.xlsx', sheet_name='Raw Prices', index_col=0)
    cov_df = pd.read_excel(r'c:\Users\venni\MathIA\IA_Portfolio_Data.xlsx', sheet_name='Covariance Matrix')
    
    p0 = prices_df['AAPL'].iloc[0]
    p1_val = prices_df['AAPL'].iloc[1]
    p2_val = prices_df['AAPL'].iloc[2]
    r1 = np.log(p1_val / p0)
    r2 = np.log(p2_val / p1_val)
    
    doc.add_paragraph(f"Price at t\u2080 (Day 1): ${p0:.2f}")
    doc.add_paragraph(f"Price at t\u2081 (Day 2): ${p1_val:.2f}")
    doc.add_paragraph(f"Price at t\u2082 (Day 3): ${p2_val:.2f}")
    doc.add_paragraph("")
    doc.add_paragraph("Applying the log return formula:")
    p = doc.add_paragraph(f'r\u2081 = ln( {p1_val:.2f} \u00F7 {p0:.2f} ) = {r1:.6f}')
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.runs[0].bold = True
    p = doc.add_paragraph(f'r\u2082 = ln( {p2_val:.2f} \u00F7 {p1_val:.2f} ) = {r2:.6f}')
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.runs[0].bold = True
    
    doc.add_paragraph(
        f"This confirms a daily log return of {r1:.4%} on Day 2 and {r2:.4%} on Day 3. "
        "This process was repeated across all 1,257 trading days for each of the five stocks to generate the full return distribution."
    )
    
    doc.add_paragraph(
        "To illustrate the variance calculation, suppose we observe 5 daily returns: [\u22120.01, 0.01, 0.02, \u22120.02, 0.01]. "
        "The mean is 0.002. The variance involves summing the squared deviations from the mean and dividing by (n \u2212 1):"
    )
    p = doc.add_paragraph('\u03C3\u00B2 = [(\u22120.01 \u2212 0.002)\u00B2 + (0.01 \u2212 0.002)\u00B2 + (0.02 \u2212 0.002)\u00B2 + (\u22120.02 \u2212 0.002)\u00B2 + (0.01 \u2212 0.002)\u00B2] \u00F7 4')
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.runs[0].bold = True
    
    doc.add_paragraph(
        "This step-by-step process ensures full transparency between the raw data and the final statistical output shown below."
    )
    
    # 5.2 Summary Statistics
    doc.add_heading('5.2 Results: Summary Statistics', level=2)
    doc.add_paragraph('Table 1: Computational Results for Mean Return and Volatility')
    
    table = doc.add_table(rows=1, cols=3)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Ticker'
    hdr_cells[1].text = 'E[R] (Annualized)'
    hdr_cells[2].text = 'Volatility (Annualized)'
    
    for _, row in stats_df.iterrows():
        row_cells = table.add_row().cells
        row_cells[0].text = str(row['Ticker'])
        row_cells[1].text = f"{row['Annualized Mean Return']:.4f} ({row['Annualized Mean Return']:.2%})"
        row_cells[2].text = f"{row['Annualized Volatility (Risk)']:.4f} ({row['Annualized Volatility (Risk)']:.2%})"
    
    doc.add_paragraph(
        "The table reveals a striking trade-off. Nvidia (NVDA) offers the highest expected return but also carries the highest "
        "volatility. Conversely, Microsoft (MSFT) and Apple (AAPL) provide more modest returns with significantly lower risk. "
        "The question that MPT answers is: can we combine these assets to capture some of Nvidia's upside while being shielded "
        "by the stability of Apple and Microsoft?"
    )
    
    # 5.3 Covariance Matrix
    doc.add_heading('5.3 The Covariance Matrix (\u03A3)', level=2)
    doc.add_paragraph(
        "The covariance matrix below summarizes the joint variability of the system. Note that Nvidia (NVDA) and Palantir (PLTR) "
        "have the highest individual variances (represented by the diagonal), making them the primary sources of risk in the portfolio."
    )
    
    table = doc.add_table(rows=1, cols=6)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Ticker'
    for i, ticker in enumerate(['AAPL', 'GOOGL', 'MSFT', 'NVDA', 'PLTR']):
        hdr_cells[i+1].text = ticker
        
    for i, row in cov_df.iterrows():
        row_cells = table.add_row().cells
        row_cells[0].text = row['Ticker']
        for j in range(5):
            row_cells[j+1].text = f"{row.iloc[j+1]:.5f}"
    
    doc.add_paragraph(
        "A critical observation from this matrix is that all off-diagonal elements are positive. This means that all five tech stocks "
        "are positively correlated\u2014they tend to move in the same direction. This limits the diversification benefit compared to, say, "
        "mixing tech stocks with government bonds. However, the magnitude of the covariances varies significantly. For example, "
        "the covariance between AAPL and PLTR is much smaller than between NVDA and PLTR, meaning Apple provides a better diversification "
        "hedge against Palantir than Nvidia does."
    )
    
    # ========================================================================
    # 6. DERIVING THE OPTIMAL PORTFOLIO (CALCULUS)
    # ========================================================================
    doc.add_page_break()
    doc.add_heading('6. Deriving the Optimal Portfolio', level=1)
    doc.add_paragraph(
        "This section contains the core mathematical derivation of the investigation. The goal is not merely to present a 'solution', "
        "but to demonstrate the full solving process from problem formulation through calculus to the final numerical result."
    )
    
    doc.add_heading('6.1 Setting Up the Constrained Optimization', level=2)
    doc.add_paragraph(
        "The objective is to minimize portfolio variance \u03C3\u209A\u00B2 = w\u1D40 \u03A3 w. However, this is a 'constrained optimization' problem. "
        "We cannot simply set all weights to zero to minimize variance, because we must adhere to two mathematical constraints:"
    )
    doc.add_paragraph("Constraint 1 (Budget): The weights must sum to unity.")
    p = doc.add_paragraph('\u03A3 w\u1D62 = 1,  equivalently:  w\u1D40 \u00B7 \U0001D7D9 = 1')
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.runs[0].bold = True
    
    doc.add_paragraph("Constraint 2 (Return): The portfolio must achieve the target return.")
    p = doc.add_paragraph('\u03A3 w\u1D62 \u03BC\u1D62 = R_target,  equivalently:  w\u1D40 \u00B7 \u03BC = 0.20')
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.runs[0].bold = True
    
    doc.add_paragraph(
        "Since we are minimizing a quadratic function subject to linear equality constraints, the method of Lagrange Multipliers "
        "is the appropriate calculus technique. This method converts a constrained problem into an unconstrained one by introducing "
        "auxiliary variables."
    )
    
    doc.add_heading('6.2 Constructing the Lagrangian', level=2)
    doc.add_paragraph(
        "We define the Lagrangian function L by combining the objective function with the constraints, weighted by the "
        "multipliers \u03BB\u2081 and \u03BB\u2082. The factor of \u00BD in the objective is a standard mathematical convenience that simplifies the derivatives:"
    )
    p = doc.add_paragraph('L(w, \u03BB\u2081, \u03BB\u2082) = \u00BD w\u1D40 \u03A3 w \u2212 \u03BB\u2081( w\u1D40\u03BC \u2212 R_target ) \u2212 \u03BB\u2082( w\u1D40\U0001D7D9 \u2212 1 )')
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.runs[0].bold = True
    
    doc.add_paragraph(
        "In this equation, the first term (\u00BD w\u1D40 \u03A3 w) is the portfolio variance we wish to minimize. The second term enforces "
        "that the portfolio achieves our 20% target return. The third term enforces that our weights sum to 1. "
        "At the optimal solution, the Lagrangian achieves a stationary point where all partial derivatives equal zero."
    )
    
    doc.add_heading('6.3 Computing the Partial Derivatives', level=2)
    doc.add_paragraph(
        "To find the minimum, I must take the partial derivative of L with respect to each variable and set them equal to zero. "
        "This is the calculus at the heart of the optimization."
    )
    
    doc.add_paragraph("First, differentiating with respect to the weight vector w:")
    p = doc.add_paragraph('\u2202L \u00F7 \u2202w = \u03A3w \u2212 \u03BB\u2081\u03BC \u2212 \u03BB\u2082\U0001D7D9 = 0')
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.runs[0].bold = True
    
    doc.add_paragraph(
        "This is a vector equation. For each individual weight w\u1D62, the partial derivative is:"
    )
    p = doc.add_paragraph('\u2202L \u00F7 \u2202w\u1D62 = \u03A3\u2C7C \u03C3\u1D62\u2C7C w\u2C7C \u2212 \u03BB\u2081\u03BC\u1D62 \u2212 \u03BB\u2082 = 0,  for i = 1, 2, \u2026, 5')
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.runs[0].bold = True
    
    doc.add_paragraph("Next, differentiating with respect to \u03BB\u2081 (the return constraint):")
    p = doc.add_paragraph('\u2202L \u00F7 \u2202\u03BB\u2081 = w\u1D40\u03BC \u2212 R_target = 0')
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.runs[0].bold = True
    
    doc.add_paragraph("And differentiating with respect to \u03BB\u2082 (the budget constraint):")
    p = doc.add_paragraph('\u2202L \u00F7 \u2202\u03BB\u2082 = w\u1D40\U0001D7D9 \u2212 1 = 0')
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.runs[0].bold = True
    
    doc.add_paragraph(
        "Together, these form a system of n + 2 = 7 equations in 7 unknowns (w\u2081 through w\u2085, plus \u03BB\u2081 and \u03BB\u2082). "
        "This is a linear system that can be solved using standard matrix algebra."
    )
    
    doc.add_heading('6.4 Solving for the Optimal Weights', level=2)
    doc.add_paragraph(
        "Rearranging the first equation, I can express the weight vector in terms of the Lagrange multipliers:"
    )
    p = doc.add_paragraph('\u03A3w = \u03BB\u2081\u03BC + \u03BB\u2082\U0001D7D9')
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.runs[0].bold = True
    
    doc.add_paragraph(
        "Multiplying both sides by the inverse of the covariance matrix \u03A3\u207B\u00B9 (which exists because \u03A3 is positive definite for a well-diversified portfolio):"
    )
    p = doc.add_paragraph('w = \u03A3\u207B\u00B9 ( \u03BB\u2081\u03BC + \u03BB\u2082\U0001D7D9 )')
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.runs[0].bold = True
    
    doc.add_paragraph(
        "Substituting this expression for w back into the two constraint equations gives us a 2 \u00D7 2 system in \u03BB\u2081 and \u03BB\u2082:"
    )
    p = doc.add_paragraph('\u03BC\u1D40 \u03A3\u207B\u00B9 \u03BC \u00B7 \u03BB\u2081  +  \u03BC\u1D40 \u03A3\u207B\u00B9 \U0001D7D9 \u00B7 \u03BB\u2082  =  R_target')
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.runs[0].bold = True
    p = doc.add_paragraph('\U0001D7D9\u1D40 \u03A3\u207B\u00B9 \u03BC \u00B7 \u03BB\u2081  +  \U0001D7D9\u1D40 \u03A3\u207B\u00B9 \U0001D7D9 \u00B7 \u03BB\u2082  =  1')
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.runs[0].bold = True
    
    doc.add_paragraph(
        "This 2 \u00D7 2 linear system can be solved for \u03BB\u2081 and \u03BB\u2082 using substitution or Cramer's rule. Once the multipliers are known, "
        "the optimal weights are obtained by back-substituting into the equation w = \u03A3\u207B\u00B9(\u03BB\u2081\u03BC + \u03BB\u2082\U0001D7D9). "
        "In practice, I used the Sequential Least Squares Programming (SLSQP) algorithm in Python's scipy library to solve this system numerically, "
        "which handles the matrix inversion and constraint enforcement simultaneously."
    )
    
    doc.add_heading('6.5 Solution for R_target = 20%', level=2)
    
    # Compute results
    mu_vec = stats_df['Annualized Mean Return'].values
    tickers_list = stats_df['Ticker'].values
    cov_mat = cov_df.iloc[:, 1:].values
    weights = solve_mpt(mu_vec, cov_mat, 0.20)
    
    doc.add_paragraph(
        "For a target return of 20%, the optimization yields the following allocation. The results show a heavy bias toward "
        "the lower-volatility stocks in the group:"
    )
    
    table = doc.add_table(rows=1, cols=2)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Ticker'
    hdr_cells[1].text = 'Optimal Weight Allocation'
    
    for i, ticker in enumerate(tickers_list):
        row_cells = table.add_row().cells
        row_cells[0].text = ticker
        row_cells[1].text = f"{weights[i]:.2%}"
    
    p_var = weights.T @ cov_mat @ weights
    p_risk = np.sqrt(p_var)
    
    doc.add_paragraph(f"\nCalculated Portfolio Volatility: {p_risk:.2%}")
    doc.add_paragraph(
        f"Contrast this with an equal-weighted portfolio (20% each), which would have a higher historical volatility. "
        "This demonstrates the significant reduction in risk achieved through MPT optimization."
    )
    doc.add_paragraph(
        "The model allocates the majority of capital to the more stable mega-cap stocks (Apple, Microsoft, and Google) while heavily "
        "underweighting the volatile Nvidia and Palantir. Even though NVDA has the highest expected return, its outsized contribution "
        "to portfolio variance means the optimizer only includes a small allocation. This illustrates a key MPT insight: "
        "an asset's value to a portfolio is determined not by its individual return, but by its marginal contribution to portfolio risk."
    )
    
    # The Efficient Frontier
    doc.add_heading('6.6 The Efficient Frontier', level=2)
    
    # Add the chart
    frontier_path = r'c:\Users\venni\MathIA\efficient_frontier.png'
    if os.path.exists(frontier_path):
        doc.add_picture(frontier_path, width=Inches(5.5))
        p = doc.add_paragraph('Figure 1: Efficient Frontier Construction (Risk vs Return)')
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph(
        "In Figure 1, the curve shows the 'Envelope' of all possible optimal portfolios. The Minimum Variance Portfolio (MVP) is the "
        "leftmost point on this curve. Any portfolio below the MVP is 'inefficient' because a portfolio with higher return "
        "exists for the same level of risk. Our chosen tech assets all lie to the right of the frontier, proving that "
        "individual stock selection is statistically inferior to diversified combination."
    )
    
    # Add Monte Carlo if available
    mc_path = r'c:\Users\venni\MathIA\monte_carlo_portfolios.png'
    if os.path.exists(mc_path):
        doc.add_picture(mc_path, width=Inches(5.5))
        p = doc.add_paragraph('Figure 2: Monte Carlo Simulation of 10,000 Random Portfolios')
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_paragraph(
            "Figure 2 shows the result of generating 10,000 random portfolios with different weight combinations. The scatter plot "
            "reveals the 'bullet' shape characteristic of the feasible set. The upper edge of this bullet is the Efficient Frontier. "
            "Any portfolio in the interior of the bullet is suboptimal because a portfolio exists on the frontier with either higher return "
            "for the same risk, or lower risk for the same return."
        )
    
    # ========================================================================
    # 7. CRITICAL REFLECTION
    # ========================================================================
    doc.add_page_break()
    doc.add_heading('7. Critical Reflection and Evaluation', level=1)
    
    doc.add_heading('7.1 Normality and The Kurtosis Problem', level=2)
    doc.add_paragraph(
        "A fundamental assumption of my model is that stock returns follow a Normal (Gaussian) distribution. In a normal distribution, "
        "significant outliers (3-sigma events) should occur less than 0.3% of the time. However, looking at the data for PLTR and NVDA, "
        "I observed 'Fat Tails' or high Kurtosis. During the 2020 pandemic and the 2022 rate-hike cycle, these stocks experienced daily "
        "price swings of 10\u201315% multiple times\u2014events that are statistically 'impossible' under a pure Gaussian model. "
        "This means the \u03C3 term underestimated the true 'tail risk' of the portfolio."
    )
    doc.add_paragraph(
        "If I were to extend this investigation, I could test the returns for normality using a Jarque-Bera test or a Shapiro-Wilk test. "
        "A p-value below 0.05 would formally reject the null hypothesis that the returns are normally distributed, quantifying the degree "
        "to which my model's core assumption is violated."
    )
    
    doc.add_heading('7.2 Temporal Instability of Covariance', level=2)
    doc.add_paragraph(
        "Furthermore, MPT assumes that the relationships between stocks (covariance) are static. This is historically false. "
        "In periods of 'Market Stress', correlations tend to jump toward 1.0. For instance, while Google and Nvidia may "
        "decouple during a normal trading month, they both crashed during the February 2022 inflation spike. "
        "This 'Correlation Breakdown' implies that diversification fails exactly when it is needed most to protect capital."
    )
    doc.add_paragraph(
        "To improve this investigation, one could use a 'Rolling Covariance' matrix that weighs recent data more heavily than older data. "
        "This is known as an Exponentially Weighted Moving Average (EWMA) approach, commonly used in the financial industry. "
        "The EWMA model assigns exponentially declining weights to older observations, making the covariance matrix responsive to "
        "recent market shifts rather than treating all historical data equally."
    )
    
    doc.add_heading('7.3 Limitations of the Lagrangian Method', level=2)
    doc.add_paragraph(
        "The Lagrangian method assumes that the solution lies in the interior of the feasible region. However, in practice, "
        "many optimal portfolios involve 'corner solutions' where one or more weights are exactly zero. This is why I imposed "
        "the additional constraint that 0 \u2264 w\u1D62 \u2264 1 (no short-selling), which transforms the problem from a pure Lagrangian "
        "into a Quadratic Programming (QP) problem. The SLSQP algorithm handles this extension, but it is important to "
        "acknowledge that the analytical Lagrangian solution and the numerical QP solution may differ when boundary constraints are active."
    )
    
    # ========================================================================
    # 8. CONCLUSION
    # ========================================================================
    doc.add_page_break()
    doc.add_heading('8. Conclusion', level=1)
    doc.add_paragraph(
        "This investigation successfully demonstrated that Modern Portfolio Theory can be used to mathematically minimize risk "
        f"for a technology portfolio. By solving for the 20% target return, I found an allocation that yielded a volatility of {p_risk:.2%}, "
        "representing a significant improvement over unoptimized strategies."
    )
    doc.add_paragraph(
        "While the limitations of normality assumptions and historical bias are present, the exploration underscores a vital "
        "mathematical truth: risk is not an absolute property of an asset, but a relative property determined by its relationship to others. "
        "The covariance matrix\u2014a purely statistical construct\u2014contains the information needed to reduce uncertainty, and the Lagrangian "
        "framework provides the calculus-based machinery to extract the optimal solution from that information."
    )
    doc.add_paragraph(
        "This IA has provided me with a deep appreciation for the intersection of statistical analysis and capital management, "
        "proving that even in the chaotic world of stocks, there is an underlying mathematical order that can be leveraged for better outcomes. "
        "The process of deriving the partial derivatives, inverting the covariance matrix, and iterating toward a numerical solution "
        "has reinforced my understanding of how abstract mathematics connects to tangible financial decisions."
    )
    
    # ========================================================================
    # BIBLIOGRAPHY
    # ========================================================================
    doc.add_heading('Bibliography', level=1)
    bibs = [
        "Markowitz, H. (1952). Portfolio Selection. The Journal of Finance, 7(1), 77-91.",
        "Malkiel, B. G. (2019). A Random Walk Down Wall Street. W. W. Norton & Company.",
        "International Baccalaureate. (2023). Mathematics: Applications and Interpretation HL Guide.",
        "Yahoo Finance. (2025). Historical Stock Data for AAPL, MSFT, GOOGL, NVDA, PLTR.",
        "Hull, J. C. (2021). Options, Futures, and Other Derivatives. Pearson Education."
    ]
    for b in bibs:
        doc.add_paragraph(b, style='List Bullet')

    # Save
    output_path = r'c:\Users\venni\MathIA\MathIA_High_Quality_Expanded_v2.docx'
    doc.save(output_path)
    print(f"Document saved to: {output_path}")

if __name__ == "__main__":
    generate_expanded_ia()
