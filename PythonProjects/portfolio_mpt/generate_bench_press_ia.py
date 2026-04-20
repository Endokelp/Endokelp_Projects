from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

def add_heading(doc, text, level):
    h = doc.add_heading(text, level=level)
    return h

def add_paragraph(doc, text, align=None, bold=False, italic=False, space_after=Pt(12)):
    p = doc.add_paragraph(text)
    p.paragraph_format.space_after = space_after
    if align:
        p.alignment = align
    if bold:
        p.runs[0].bold = True
    if italic:
        p.runs[0].italic = True
    return p

def generate_bench_press_ia():
    doc = Document()
    
    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Inter'
    font.size = Pt(11)
    
    # --- Title Page ---
    doc.add_heading('Mathematics: Applications and Interpretation HL', 0)
    doc.add_heading('Internal Assessment', 1)
    doc.add_paragraph('\n\n\n\n')
    title = doc.add_paragraph('The Impact of Auditory Stimuli on Upper Body Power and Muscular Endurance: A Statistical Correlation Analysis')
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.runs[0].bold = True
    title.runs[0].font.size = Pt(22)
    
    doc.add_paragraph('\n\n\n')
    q = doc.add_paragraph('Research Question: Is there a significant correlation in individual’s bench press strength and endurance when listening to music compared to lifting in silence?')
    q.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph('\n\n\n\n\n')
    doc.add_paragraph('Candidate Session Number: [REDACTED]')
    doc.add_paragraph('Word Count: ~1650')
    doc.add_page_break()
    
    # --- 1. Introduction ---
    add_heading(doc, '1. Introduction', 1)
    add_heading(doc, '1.1 Personal Engagement and Rationale', 2)
    intro_text = (
        "Walking into a gym, one of the most consistent sights is the presence of headphones. "
        "Whether it is heavy metal, upbeat pop, or aggressive rap, music seems inseparable from the lifting experience. As a frequent gym-goer, "
        "I have often wondered if my personal 'PRs' (Personal Records) are a product of my training or if my Spotify playlist is doing the heavy lifting. "
        "I have noticed that on days when I forget my headphones, the weights feel heavier and my sets feel shorter. This led me to a curiosity: "
        "is the effect of music merely psychological, or does it manifest as a statistically significant increase in physical output?\n\n"
        "This fascination goes beyond mere observation. Throughout my high school career, I have competed in local powerlifting meets, where "
        "the atmosphere is carefully curated with high-intensity soundscapes designed to elicit maximal physical response. However, 'feeling' "
        "stronger is not the same as 'being' stronger in a mathematically verifiable way. The human brain is a master of placebo; I want to "
        "strip away that subjectivity. This investigation allows me to combine my passion for strength training with the rigors of statistical "
        "analysis. While many studies explore how music affects cardiovascular performance (like running or cycling), there is less emphasis on the "
        "correlation between absolute strength (1-Rep Max) and relative endurance to failure within the same individual under different auditory conditions. "
        "By applying Pearson’s correlation, paired t-tests, and the Fisher Z-transformation, I aim to uncover the mathematical relationship between "
        "sound and strength."
    )
    doc.add_paragraph(intro_text)
    
    add_heading(doc, '1.2 Aim and Research Question', 2)
    aim_text = (
        "The aim of this investigation is to determine if music acts as a performance enhancer for bench press strength and endurance, and "
        "critically, if it alters the relationship between these two metrics. The Research Question is: Is there a significant correlation in "
        "individual’s bench press strength and endurance when listening to music compared to lifting in silence?\n\n"
        "To address this, I will evaluate the strength-endurance relationship using regression analysis and determine if the difference in "
        "performance across conditions is statistically significant using inferential testing."
    )
    doc.add_paragraph(aim_text)
    
    # --- 2. Mathematical Background ---
    add_heading(doc, '2. Mathematical Background', 1)
    
    add_heading(doc, '2.1 Descriptive Statistics', 2)
    m1 = (
        "To understand the central tendency and dispersion of the data, I will calculate the Mean (μ) and Standard Deviation (σ). "
        "The standard deviation is particularly important as it indicates the consistency of the participants' performance. A high σ "
        "would suggest that the effects of music vary wildly between individuals, whereas a low σ would suggest a more uniform response. "
        "The formula for the sample standard deviation is:\n"
    )
    doc.add_paragraph(m1)
    p = doc.add_paragraph('σ = √[ Σ (xᵢ - x̄)² / (n - 1) ]')
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    add_heading(doc, '2.2 Pearson Correlation Coefficient (r)', 2)
    m2 = (
        "Pearson’s r measures the strength and direction of a linear relationship between two variables—in this case, 1RM strength (X) and "
        "repetitions to failure at 70% intensity (Y). One of the assumptions of the Pearson r is that the variables are normally distributed. "
        "The coefficient is calculated by dividing the covariance of the two variables by the product of their standard deviations:\n"
    )
    doc.add_paragraph(m2)
    p = doc.add_paragraph('r = Σ((xᵢ - x̄)(yᵢ - ȳ)) / √(Σ(xᵢ - x̄)² Σ(yᵢ - ȳ)²)')
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph("\nA value of r = 0 indicates no linear relationship, while r = 1 indicates a perfect positive correlation. In the context of "
                      "strength training, we might expect a positive r because participants with a higher absolute strength (1RM) might "
                      "possess more muscle mass, which could theoretically allow for more efficient submaximal repetition performance.")
    
    add_heading(doc, '2.3 Fisher Z-Transformation', 2)
    m_fisher = (
        "To compare two correlation coefficients (r_music vs r_silence), we cannot simply compare their raw values because the sampling "
        "distribution of r is not normal, especially when |r| is high. The Fisher Z-transformation converts r into a value Z' that "
        "follows a normal distribution:\n"
    )
    doc.add_paragraph(m_fisher)
    p = doc.add_paragraph("Z' = 0.5 * ln( (1 + r) / (1 - r) )")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph("\nOnce transformed, we can use the standard error of the difference to calculate a z-score and determine if the "
                      "difference between the two correlations is statistically significant. The Standard Error (SE) for the difference "
                      "between two Z' values is given by:\n")
    p = doc.add_paragraph("SE = √[ 1 / (n₁ - 3) + 1 / (n₂ - 3) ]")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    add_heading(doc, '2.4 Paired T-Test', 2)
    m3 = (
        "Since the same participants were tested in both 'Music' and 'Silence' conditions, a paired t-test is appropriate. This test evaluates "
        "if the mean difference between the two conditions is significantly different from zero, accounting for the inherent variability of "
        "each participant. This 'within-subjects' design is powerful because it controls for individual differences in strength baseline. "
        "The t-statistic is calculated as:\n"
    )
    doc.add_paragraph(m3)
    p = doc.add_paragraph('t = d̄ / (s_d / √n)')
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph("where d̄ is the mean difference, s_d is the standard deviation of the differences, and n is the sample size (n=30).")
    
    # --- 3. Methodology ---
    add_heading(doc, '3. Methodology and Data Collection', 1)
    meth_text = (
        "This investigation used a sample of 30 male participants (ages 18-25) with at least one year of consistent weightlifting experience. "
        "This experience level was chosen to ensure that 'newbie gains' or learning effects—where performance improves simply because the participant "
        "is getting better at the movement—would be minimized. \n\n"
        "The experimental design involved two sessions separated by 48 hours to ensure full muscular recovery. To prevent an 'order effect' "
        "(where Participants might perform better in the second session simply because they were more familiar with the lab environment), "
        "the order of music and silence was counterbalanced: half started with music, half with silence.\n\n"
        "In Session 1, participants performed a 1RM bench press test using a standardized protocol of progressive loading. After achieving their 1RM, "
        "a 15-minute rest period was provided where they remained seated to lower heart rate to baseline. They then performed a single set "
        "to failure at exactly 70% of their 1RM. In Session 2, the process was repeated, but participants listened to high-tempo music (130-150 BPM). "
        "Crucially, the 70% load in Session 2 was kept identical to the absolute weight used in Session 1, rather than being recalculated from their "
        "Session 2 1RM. This ensured that 'Endurance' was measured against a fixed resistance relative to their baseline strength."
    )
    doc.add_paragraph(meth_text)
    
    # --- 4. Data Analysis ---
    add_heading(doc, '4. Data Analysis', 1)
    
    add_heading(doc, '4.1 Raw Data Presentation', 2)
    doc.add_paragraph("The following table summarizes the performance metrics for a subset of participants (1-15).")
    
    # Data Data
    strength_no = [160, 155, 105, 135, 205, 185, 90, 210, 110, 115, 125, 95, 105, 180, 225, 145, 100, 90, 95, 120, 140, 155, 105, 195, 95, 130, 115, 120, 85, 65]
    strength_yes = [165, 155, 110, 130, 205, 190, 95, 220, 110, 110, 125, 95, 105, 185, 225, 150, 105, 90, 95, 115, 140, 155, 105, 195, 100, 135, 120, 120, 80, 70]
    endurance_no = [12, 12, 9, 14, 11, 10, 9, 12, 14, 8, 10, 11, 13, 9, 10, 14, 12, 14, 9, 10, 11, 12, 9, 11, 10, 12, 13, 10, 8, 11]
    endurance_yes = [14, 12, 11, 15, 14, 12, 11, 12, 17, 10, 13, 13, 15, 11, 12, 16, 14, 17, 11, 12, 13, 15, 11, 12, 12, 14, 16, 11, 10, 12]

    table = doc.add_table(rows=1, cols=5)
    table.style = 'Table Grid'
    hdr = table.rows[0].cells
    hdr[0].text = 'ID'
    hdr[1].text = '1RM Silence'
    hdr[2].text = '1RM Music'
    hdr[3].text = 'Endur Silence'
    hdr[4].text = 'Endur Music'
    
    for i in range(15): 
        row = table.add_row().cells
        row[0].text = str(i+1)
        row[1].text = str(strength_no[i])
        row[2].text = str(strength_yes[i])
        row[3].text = str(endurance_no[i])
        row[4].text = str(endurance_yes[i])
    
    doc.add_paragraph("\nTable 1: Partial display of the performance dataset. The full dataset of 30 participants was used for all calculations.")
    
    add_heading(doc, '4.2 Descriptive Statistics and Comparison', 2)
    stats_text = (
        "The summary statistics reveal a clear trend: both strength and endurance increased when music was introduced. "
        "The mean 1RM increased from 131.83 lbs to 133.33 lbs, while the mean repetitions increased significantly from 11.00 to 12.93.\n\n"
        "Interestingly, the standard deviation for endurance also increased from 1.80 to 2.00. This suggests that while music "
        "increases performance on average, individuals respond to it with different levels of intensity. Some participants "
        "showed a gain of 3-4 repetitions with music, while others remained static."
    )
    doc.add_paragraph(stats_text)
    
    # Summary Table
    table = doc.add_table(rows=5, cols=3)
    table.style = 'Table Grid'
    table.rows[0].cells[0].text = 'Metric'
    table.rows[0].cells[1].text = 'Silence (Mean ± SD)'
    table.rows[0].cells[2].text = 'Music (Mean ± SD)'
    
    table.rows[1].cells[0].text = '1RM Strength (lbs)'
    table.rows[1].cells[1].text = '131.83 ± 41.59'
    table.rows[1].cells[2].text = '133.33 ± 42.37'
    
    table.rows[2].cells[0].text = 'Endurance (Reps)'
    table.rows[2].cells[1].text = '11.00 ± 1.80'
    table.rows[2].cells[2].text = '12.93 ± 2.00'
    
    table.rows[3].cells[0].text = 'T-Test Significance'
    table.rows[3].cells[1].text = 'Baseline'
    table.rows[3].cells[2].text = 'p=0.036* | p<0.001***'
    
    table.rows[4].cells[0].text = 'Correlation (r)'
    table.rows[4].cells[1].text = '0.0783'
    table.rows[4].cells[2].text = '0.0027'
    
    doc.add_paragraph("\nTable 2: Statistical Summary Table")
    
    add_heading(doc, '4.3 Normal Distribution and Skewness', 2)
    norm_text = (
        "One of the key assumptions of Pearson’s r and the paired t-test is that the differences between the datasets are normally distributed. "
        "I analyzed the distribution of the repetition gains (Music reps - Silence reps). The skewness of this distribution was found to be "
        "0.24, which is within the acceptable range for normality (-1 to +1). This indicates that the parametric tests used in this "
        "investigation are mathematically sound."
    )
    doc.add_paragraph(norm_text)
    
    add_heading(doc, '4.4 Correlation Analysis: Strength vs Endurance', 2)
    corr_body = (
        "The most surprising result of this investigation is the lack of correlation between absolute 1RM strength and relative endurance. "
        "In the silence condition, the Pearson correlation coefficient was r = 0.0783. This is an extremely weak positive correlation, "
        "accounting for only ~0.6% of the shared variance (r² = 0.0061).\n\n"
        "When music was introduced, the correlation plummeted even further to r = 0.0027. This suggests that the presence of music "
        "acts as a 'randomizing' agent in the relationship between these two physical traits. To verify if this drop in correlation "
        "was statistically significant, I performed a Fisher Z-transformation. The calculated z-score for the difference between "
        "the two correlations was z = -0.278, with a corresponding p-value of 0.781. Since p > 0.05, we fail to reject the null hypothesis; "
        "the change in correlation is not statistically significant and is likely due to sampling variance."
    )
    doc.add_paragraph(corr_body)
    
    # Plot 1
    if os.path.exists(r'c:\Users\venni\MathIA\strength_endurance_silence.png'):
        doc.add_picture(r'c:\Users\venni\MathIA\strength_endurance_silence.png', width=Inches(4.5))
        p = doc.add_paragraph('Figure 1: Scatter Plot of Strength vs Endurance (Silence Condition)')
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # --- 5. Critical Reflection ---
    add_heading(doc, '5. Evaluation and Reflection', 1)
    
    add_heading(doc, '5.1 Significance of the P-Values', 2)
    ref1 = (
        "The paired t-test for endurance yielded a p-value of less than 0.001, which is highly significant. This results aligns with the "
        "'Psychogenic Arousal' theory, which posits that high-tempo music increases heart rate and cortisol levels, allowing for "
        "delayed perception of effort. In contrast, the 1RM strength gain (p = 0.0366) was less pronounced. This makes sense from a physiological "
        "standpoint: maximal force (1RM) is limited by motor unit recruitment and cross-bridge formation, which are hard biological limits. "
        "Endurance, however, is heavily influenced by pain tolerance and mental fatigue, areas where music is known to be most effective."
    )
    doc.add_paragraph(ref1)
    
    add_heading(doc, '5.2 Limitations and Sources of Error', 2)
    ref2 = (
        "A major limitation of this study is 'Participant Habituation.' Many modern lifters have never trained in silence. This lack of "
        "familiarity might have induced a 'frustration effect' or lowered morale in the silence condition, creating an artificial performance "
        "gap. A better design might include a 'Control' group that listens to neutral white noise.\n\n"
        "The 'Specific Genre' effect was not explored. It is possible that participants who chose aggressive genres (like Metal) "
        "had a higher surge in adrenaline compared to those who chose upbeat Pop, despite similar BPM. This variation introduces "
        "uncontrolled noise into the dataset."
    )
    doc.add_paragraph(ref2)
    
    # --- 6. Conclusion ---
    add_heading(doc, '6. Conclusion', 1)
    conc_text = (
        "This investigation set out to determine if music alters the correlation between bench press strength and endurance. "
        "The results demonstrate that while music significantly improves both metrics in isolation—increasing 1RM by 1.1% "
        "and endurance by 17.5%—it does not create a linear relationship between them. Strength and endurance remain "
        "statistically independent regardless of the auditory environment.\n\n"
        "For the athlete, this means that while music is a powerful ergogenic aid, it cannot transform a 'strength-focused' lifter "
        "into an 'endurance-focused' one. There is no underlying mathematical bridge between maximal force and repetition stability "
        "that music can manifest. This IA has highlighted the complexity of human performance and the utility of "
        "statistics in separating genuine physiological trends from the noise of the gym floor."
    )
    doc.add_paragraph(conc_text)
    
    # --- Bibliography ---
    add_heading(doc, 'Bibliography', 1)
    bibs = [
        "Karageorghis, C. I., & Priest, D. L. (2012). Music in the exercise domain: a review and synthesis (Part I). International Review of Sport and Exercise Psychology.",
        "American College of Sports Medicine. (2018). Guidelines for Exercise Testing and Prescription. Wolters Kluwer.",
        "International Baccalaureate. (2023). Mathematics: Applications and Interpretation HL Guide.",
        "StatsDirect. (2024). Pearson's Correlation and Fisher Z-Transformation Implementation.",
        "Yerkes, R. M., & Dodson, J. D. (1908). The relation of strength of stimulus to rapidity of habit-formation. Journal of Comparative Neurology and Psychology."
    ]
    for b in bibs:
        doc.add_paragraph(b, style='List Bullet')
    
    output_path = r'c:\Users\venni\MathIA\MathIA_Bench_Press_Final.docx'
    doc.save(output_path)
    print(f"Document saved to: {output_path}")

if __name__ == "__main__":
    generate_bench_press_ia()
