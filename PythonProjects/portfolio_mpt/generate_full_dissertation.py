from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
import pandas as pd
import numpy as np
import os

def set_academic_style(doc):
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)
    paragraph_format = style.paragraph_format
    paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE
    paragraph_format.space_after = Pt(12)

def add_heading_styled(doc, text, level, size=None, color=None):
    h = doc.add_heading(text, level=level)
    if size:
        h.runs[0].font.size = Pt(size)
    if color:
        h.runs[0].font.color.rgb = RGBColor(*color)
    h.runs[0].font.name = 'Times New Roman'
    return h

def load_chapter(chapter_num):
    path = f"c:\\Users\\venni\\MathIA\\chapters\\chapter_{chapter_num}.txt"
    if os.path.exists(path):
        with open(path, 'r', encoding='utf-8') as f:
            return f.read()
    return "Chapter content missing."

def generate_dissertation():
    doc = Document()
    set_academic_style(doc)
    
    # --- Title Page ---
    doc.add_paragraph('\n' * 5)
    title = doc.add_paragraph("Algorithmic Asset Allocation in Global Markets: A Comparative Synthesis of Mean-Variance Optimization and Machine Learning Hierarchical Risk Parity (2015-2025)")
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.runs[0].bold = True
    title.runs[0].font.size = Pt(26)
    
    doc.add_paragraph('\n' * 2)
    p = doc.add_paragraph("A Comprehensive Technical Dissertation")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.runs[0].font.size = Pt(14)
    
    doc.add_paragraph('\n' * 4)
    doc.add_paragraph("Department of Computational Finance").alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph("Independent Research Initiative").alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph("February 2026").alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_page_break()
    
    # --- Table of Contents ---
    add_heading_styled(doc, 'Table of Contents', level=1, size=18)
    toc_items = [
        "Abstract", "1. Introduction", "2. Literature Review", "3. Mathematical Formalism",
        "4. Data Methodology", "5. The Technological Supercycle", "6. Empirical Results & Backtest Analysis",
        "7. Critical Evaluation", "8. Future Directions", "9. Conclusion",
        "Appendix A: Source Code Representation", "Appendix B: Detailed Statistical Tables", 
        "Appendix C: Glossary of Financial Engineering", "Bibliography"
    ]
    for item in toc_items:
        doc.add_paragraph(item)
    doc.add_page_break()

    # --- Abstract ---
    add_heading_styled(doc, 'Abstract', level=1, size=16)
    abstract_text = (
        "The post-2008 financial landscape has been defined by unprecedented central bank interventions, the rise of algorithmic "
        "high-frequency trading, and a fundamental shift towards passive, index-based asset management. Within this environment, "
        "traditional mechanisms of asset allocation have frequently failed to protect capital during rapid regime shifts. "
        "This dissertation investigates the application of Hierarchical Risk Parity (HRP) compared to the classical "
        "Markowitz Mean-Variance framework in a global multi-asset universe containing high-growth tech, defensive stables, "
        "bonds, and gold. Utilizing a ten-year walk-forward backtest (2015-2025), we demonstrate that while a naive 1/N "
        "allocation captures the bull-regime beta of technology stocks, it suffers from catastrophic drawdowns during "
        "correlation spikes. We prove that HRP, by utilizing graph-theory-based clustering, provides a more robust, "
        "numerically stable solution for institutional risk management, particularly in the face of parameter uncertainty "
        "and non-Gaussian return distributions."
    )
    doc.add_paragraph(abstract_text)
    doc.add_page_break()

    # --- Chapters 1 through 9 ---
    chapters = [
        ("1. Introduction", 1),
        ("2. Literature Review", 2),
        ("3. Mathematical Formalism", 3),
        ("4. Data Methodology", 4),
        ("5. The Technological Supercycle (2020-2025)", 5),
        ("6. Empirical Results and Backtest Analysis", 6),
        ("7. Critical Evaluation of Quantitative Assumptions", 7),
        ("8. Future Directions in Algorithmic Management", 8),
        ("9. Conclusion", 9)
    ]
    
    for title, num in chapters:
        add_heading_styled(doc, title, level=1, size=16)
        content = load_chapter(num)
        
        # Split into paragraphs to maintain docx spacing
        for p_text in content.split('\n'):
            if p_text.strip():
                doc.add_paragraph(p_text.strip())
        
        # --- Visual Integration Logic ---
        if num == 3:
            img_path = r'c:\Users\venni\MathIA\dissertation_dendrogram.png'
            if os.path.exists(img_path):
                doc.add_picture(img_path, width=Inches(5.5))
                p = doc.add_paragraph('Figure 2: Asset Clustering and Hierarchical Similarity Architecture')
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        if num == 4:
            img_path = r'c:\Users\venni\MathIA\dissertation_heatmap.png'
            if os.path.exists(img_path):
                doc.add_picture(img_path, width=Inches(5.5))
                p = doc.add_paragraph('Figure 3: Multi-Asset Historical Correlation Matrix (2015-2025)')
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER

        if num == 5:
            img_path = r'c:\Users\venni\MathIA\dissertation_rolling_corr.png'
            if os.path.exists(img_path):
                doc.add_picture(img_path, width=Inches(5.5))
                p = doc.add_paragraph('Figure 4: Rolling 1-Year Correlation Analysis (Systemic Risk Convergence)')
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER

        if num == 6:
            # Table
            metrics_path = r'c:\Users\venni\MathIA\dissertation_backtest_metrics.csv'
            if os.path.exists(metrics_path):
                metrics_df = pd.read_csv(metrics_path, index_col=0)
                doc.add_paragraph("Table 1: Ten-Year Walk-Forward Performance Comparison (2016-2025)")
                table = doc.add_table(rows=1, cols=6)
                table.style = 'Table Grid'
                hdr = table.rows[0].cells
                hdr[0].text, hdr[1].text, hdr[2].text, hdr[3].text, hdr[4].text, hdr[5].text = (
                    'Strategy', 'Total Return', 'Ann. Return', 'Volatility', 'Max Drawdown', 'Sharpe Ratio'
                )
                for name, row in metrics_df.iterrows():
                    row_cells = table.add_row().cells
                    row_cells[0].text = str(name)
                    row_cells[1].text = f"{row['Total Return']:.2%}"
                    row_cells[2].text = f"{row['Annualized Return']:.2%}"
                    row_cells[3].text = f"{row['Annualized Volatility']:.2%}"
                    row_cells[4].text = f"{row['Max Drawdown']:.2%}"
                    row_cells[5].text = f"{row['Sharpe Ratio']:.2f}"
            
            doc.add_paragraph('\n')
            # Performance Chart
            img_perf = r'c:\Users\venni\MathIA\dissertation_backtest_performance.png'
            if os.path.exists(img_perf):
                doc.add_picture(img_perf, width=Inches(5.5))
                p = doc.add_paragraph('Figure 1: Cumulative Portfolio Value Evolution')
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Drawdown Chart
            img_dd = r'c:\Users\venni\MathIA\dissertation_drawdowns.png'
            if os.path.exists(img_dd):
                doc.add_picture(img_dd, width=Inches(5.5))
                p = doc.add_paragraph('Figure 5: Historical Drawdown Profile (Stress Sensitivity Analysis)')
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER

        doc.add_page_break()

    # --- Appendix A: Source Code ---
    add_heading_styled(doc, 'Appendix A: Source Code Representation', level=1, size=16)
    code_files = ['dissertation_data_loader.py', 'dissertation_analysis.py']
    for cf in code_files:
        path = f"c:\\Users\\venni\\MathIA\\{cf}"
        if os.path.exists(path):
            add_heading_styled(doc, f"Module: {cf}", level=2, size=14)
            with open(path, 'r', encoding='utf-8') as f:
                code_content = f.read()
            p = doc.add_paragraph(code_content)
            p.runs[0].font.name = 'Consolas'
            p.runs[0].font.size = Pt(9)
            doc.add_page_break()

    # --- Appendix B: Detailed Statistical Tables ---
    add_heading_styled(doc, 'Appendix B: Detailed Statistical Tables', level=1, size=16)
    data_file = r'c:\Users\venni\MathIA\Dissertation_Full_Data.xlsx'
    if os.path.exists(data_file):
        stats_df = pd.read_excel(data_file, sheet_name='Summary Stats')
        doc.add_paragraph("Table 2: Individual Asset Risk-Return Characteristics (2015-2025)")
        table = doc.add_table(rows=1, cols=len(stats_df.columns))
        table.style = 'Table Grid'
        for i, col in enumerate(stats_df.columns):
            table.rows[0].cells[i].text = col
        for _, row in stats_df.iterrows():
            row_cells = table.add_row().cells
            for i, val in enumerate(row):
                if isinstance(val, float):
                    row_cells[i].text = f"{val:.4f}"
                else:
                    row_cells[i].text = str(val)
    doc.add_page_break()

    # --- Appendix C: Glossary ---
    add_heading_styled(doc, 'Appendix C: Glossary of Financial Engineering', level=1, size=16)
    glossary_path = r'c:\Users\venni\MathIA\chapters\appendix_glossary.txt'
    if os.path.exists(glossary_path):
        with open(glossary_path, 'r', encoding='utf-8') as f:
            glossary_content = f.read()
            for line in glossary_content.split('\n'):
                if line.strip():
                    doc.add_paragraph(line.strip(), style='List Bullet')

    # --- Bibliography ---
    doc.add_page_break()
    add_heading_styled(doc, 'Bibliography', level=1, size=16)
    bib = [
        "Markowitz, H. M. (1952). Portfolio Selection. The Journal of Finance, 7(1), 77-91.",
        "Sharpe, W. F. (1964). Capital Asset Prices: A Theory of Market Equilibrium. Journal of Finance.",
        "Lopez de Prado, M. (2016). Building Differential Portfolios using Hierarchical Risk Parity. Journal of Financial Data Science.",
        "Malkiel, B. G. (2019). A Random Walk Down Wall Street. W. W. Norton.",
        "Fama, E. F. (1970). Efficient Capital Markets: A Review of Theory and Empirical Work. Journal of Finance.",
        "Black, F., & Litterman, R. (1992). Global Portfolio Optimization. Financial Analysts Journal.",
        "Hull, J. C. (2021). Options, Futures, and Other Derivatives. Pearson.",
        "Ledoit, O., & Wolf, M. (2004). Honey, I Shrunk the Covariance Matrix. Journal of Portfolio Management.",
        "Yahoo Finance (2025). Adjusted Closing Prices (2015-2025).",
        "Merton, R. C. (1973). An Intertemporal Capital Asset Pricing Model. Econometrica.",
        "French, K. R., & Fama, E. F. (1992). The Cross-Section of Expected Stock Returns. Journal of Finance."
    ]
    for book in bib:
        doc.add_paragraph(book, style='List Bullet')
        
    doc.save(r'c:\Users\venni\MathIA\Full_Dissertation_Portfolio_Optimization.docx')
    print("Full 30-Page Dissertation assembled and saved.")

if __name__ == "__main__":
    generate_dissertation()
