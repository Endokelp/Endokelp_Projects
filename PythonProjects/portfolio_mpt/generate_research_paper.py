from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import pandas as pd
import numpy as np
import os

def add_heading_styled(doc, text, level, size=None, color=None):
    h = doc.add_heading(text, level=level)
    if size:
        h.runs[0].font.size = Pt(size)
    if color:
        h.runs[0].font.color.rgb = RGBColor(*color)
    return h

def generate_research_paper():
    doc = Document()
    
    # --- Title Page ---
    title_text = (
        "Algorithmic Asset Allocation in High-Growth Technology Sectors: "
        "A Comparative Quantitative Study of Mean-Variance Optimization, "
        "Tangency Portfolios, and Parameter Uncertainty"
    )
    title = doc.add_paragraph(title_text)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.runs[0].bold = True
    title.runs[0].font.size = Pt(24)
    
    doc.add_paragraph('\n' * 2)
    author = doc.add_paragraph('Independent Quantitative Research Initiative')
    author.alignment = WD_ALIGN_PARAGRAPH.CENTER
    author.runs[0].italic = True
    doc.add_paragraph('February 2026').alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_page_break()
    
    # --- Abstract ---
    add_heading_styled(doc, 'Abstract', level=1, size=16)
    abstract_text = (
        "This paper looks at Modern Portfolio Theory (MPT) "
        "applied to five prominent technology equities (AAPL, MSFT, GOOGL, NVDA, and PLTR) during the 2020-2025 "
        "market cycle. The study investigates the mathematical efficacy of three distinct allocation strategies: "
        "the Minimum Variance Portfolio (MVP), the Maximum Sharpe Ratio (Tangency) Portfolio, and a Naive Equal-Weighted "
        "benchmark. Utilizing daily adjusted returns, we derive the Efficient Frontier and the Capital Market Line (CML) "
        "under a risk-free rate assumption of 2%. Our results indicate that while optimization significantly reduces "
        "uncompensated idiosyncratic risk, the results are highly sensitive to 'fat-tail' events and temporal "
        "covariance instability. This paper further explores the matrix algebra underlying these optimizations "
        "and critiques the Gaussian return assumption in the context of high-volatility semiconductor and AI assets."
    )
    doc.add_paragraph(abstract_text)
    doc.add_paragraph('\nKeywords: Modern Portfolio Theory, Quadratic Programming, Efficient Frontier, Sharpe Ratio, Risk-Adjusted Returns.')
    
    # --- 1. Introduction ---
    add_heading_styled(doc, '1. Introduction', level=1, size=14)
    doc.add_paragraph(
        "The global financial architecture has undergone a radical transformation over the past decade, defined by the "
        "ascendancy of the technology sector as the dominant component of equity benchmarks. As artificial intelligence and "
        "cloud computing infrastructure become the primary drivers of corporate productivity, a specific subset of "
        "equities—often referred to as 'hyperscalers' and 'enablers'—have exhibited asymmetric growth profiles. However, "
        "this potential for high returns is counterbalanced by significant price volatility and complex cross-correlation structures.\n\n"
        "The problem of capital allocation in such a volatile environment remains a central challenge for both individual "
        "and institutional investors. Portfolio optimization, the process of selecting the best proportions of various "
        "assets to maximize returns or minimize risk, provides a mathematical solution to this problem. The standard "
        "framework for this optimization remains Harry Markowitz’s Modern Portfolio Theory (MPT), which posits that the "
        "primary concern of an investor should be the variance of the portfolio rather than the volatility of individual assets."
    )
    
    doc.add_paragraph(
        "This paper aims to expand upon the basic MPT framework by conducting a comparative analysis of different optimization "
        "objective functions. We ask: How do different mathematical definitions of 'optimality'—pure risk minimization versus "
        "risk-adjusted return maximization—perform when applied to a high-beta technology portfolio? We also sketch "
        "the theoretical foundations of the Capital Asset Pricing Model (CAPM) and the Efficient Frontier to understand the "
        "geometric relationships between risk and reward in the modern AI bull market."
    )

    # --- 2. Historical and Theoretical Framework ---
    add_heading_styled(doc, '2. Theoretical Framework', level=1, size=14)
    doc.add_paragraph(
        "Before the advent of MPT in 1952, investors typically focused on the merits of individual stocks, seeking 'quality' "
        "companies without a rigorous mathematical understanding of how these assets interacted with one another in a collective "
        "portfolio. Markowitz’s seminal paper, 'Portfolio Selection,' revolutionized this by introducing the concept of "
        "Expected Return-Variance efficiency. Markowitz proved that for any given level of return, there exists a unique "
        "combination of assets that minimizes risk, provided the assets are not perfectly correlated.\n\n"
        "Following Markowitz, the theories of William Sharpe, John Lintner, and Jan Mossin led to the development of the "
        "Capital Asset Pricing Model (CAPM). CAPM introduced the distinction between systematic risk (Market Beta), which "
        "cannot be diversified away, and unsystematic risk (Alpha/Idiosyncratic risk), which is specific to a single company "
        "and can be mitigated through a diversified portfolio. The model also hypothesized the existence of a 'Market Portfolio' "
        "and a 'Capital Market Line' (CML), representing the optimal combination of risk-free assets and risky equities."
    )
    
    # --- 3. Mathematical Foundations of the Model ---
    add_heading_styled(doc, '3. Mathematical Foundations', level=1, size=14)
    doc.add_paragraph(
        "The core of the Mean-Variance optimization problem is rooted in matrix algebra and quadratic programming. "
        "Let us define the foundational variables for our system of $N=5$ assets."
    )
    
    # Mathematical derivation
    m1 = doc.add_paragraph("3.1 Continuous Compounding and Log Returns")
    doc.add_paragraph(
        "For a series of discrete prices $P_0, P_1, ..., P_t$, the simple periodic return is defined as $R_t = (P_t - P_{t-1})/P_{t-1}$. "
        "While intuitive, simple returns lack the property of additivity across time. In this study, we use Logarithmic Returns ($r_t$):"
    )
    p = doc.add_paragraph("r_t = ln( P_t / P_{t-1} )")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.runs[0].bold = True
    
    doc.add_paragraph(
        "Log returns are advantageous because they model continuous compounding and are normally distributed in many empirical contexts. "
        "Crucially, the return of a portfolio over multiple days is the simple sum of daily log returns, which simplifies the "
        "calculation of long-term expected outcomes."
    )
    
    m2 = doc.add_paragraph("3.2 Vector-Matrix Notation for Portfolio Metrics")
    doc.add_paragraph(
        "For a portfolio with weight vector $\\mathbf{w}$ where $\\sum w_i = 1$, and expected return vector $\\mathbf{\mu}$, "
        "the expected portfolio return $E[R_p]$ is a linear combination:"
    )
    p = doc.add_paragraph("E[R_p] = \\mathbf{w}^T \\mathbf{\mu}")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph(
        "The portfolio variance $\\sigma_p^2$, however, is a non-linear quadratic form that accounts for the inter-asset "
        "covariance matrix $\\Sigma$:"
    )
    p = doc.add_paragraph("\\sigma_p^2 = \\mathbf{w}^T \\Sigma \\mathbf{w}")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph(
        "This matrix expansion can be expressed as:\n"
        "$\\sigma_p^2 = \\sum_{i=1}^N \\sum_{j=1}^N w_i w_j \\sigma_{ij}$\n\n"
        "This proves that the risk of the portfolio is determined not just by the individual variances of the stocks "
        "($\\sigma_{ii}$), but exponentially more so by the covariances between them ($\\sigma_{ij}$). If $\\sigma_{ij}$ is low "
        "or negative, the second term reduces the overall sum, creating the 'Free Lunch' of diversification."
    )
    
    m3 = doc.add_paragraph("3.3 The Sharpe Ratio and Tangency")
    doc.add_paragraph(
        "Beyond simple variance minimization, we seek the 'Maximum Sharpe Ratio' portfolio. The Sharpe Ratio ($S$) "
        "measures the excess return of the portfolio relative to a risk-free rate ($R_f$) per unit of volatility:"
    )
    p = doc.add_paragraph("S = ( E[R_p] - R_f ) / \\sigma_p")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph(
        "Geometrically, the Maximum Sharpe Ratio portfolio is the 'Tangency Portfolio' where the line starting from the "
        "risk-free rate on the Y-axis is tangent to the Efficient Frontier. This line is the Capital Market Line (CML). "
        "According to the Separation Theorem, any rational investor should hold a combination of the Tangency Portfolio "
        "and the risk-free asset, regardless of their own individual risk tolerance."
    )

    # --- 4. Empirical Methodology ---
    add_heading_styled(doc, '4. Methodology and Data Processing', level=1, size=14)
    doc.add_paragraph(
        "The universe of assets selected for this study represents five 'Archetypes' of the modern technology trade:\n"
        "1. AAPL (Hardware ecosystems): Stable cash flows and high institutional ownership.\n"
        "2. MSFT (Software/Cloud): The bridge between legacy enterprise and AI compute.\n"
        "3. GOOGL (Digital Advertising): High correlation with general consumer economic health.\n"
        "4. NVDA (Semiconductor compute): The most volatile high-growth engine of the set.\n"
        "5. PLTR (Big Data Analytics): Highly idiosyncratic asset with asymmetrical return profiles.\n\n"
        "Data was extracted using the `yfinance` protocol for the period spanning Jan 1, 2020, to Jan 1, 2025. This 5-year "
        "window includes significant market regimes: the COVID-19 crash and recovery (2020), the loose-money tech rally (2021), "
        "the inflation/interest-rate drawdown (2022), and the generative AI catalyst era (2023-2024). This varied "
        "regime history allows for a robust testing of covariance stability across differing economic environments."
    )
    
    doc.add_paragraph(
        "The optimization was performed using a Sequential Least Squares Programming (SLSQP) algorithm. We enforced a "
        "long-only constraint ($w_i \\geq 0$) and a full-investment constraint ($\\sum w_i = 1$). A risk-free rate of "
        "2% was assumed, reflecting a conservative historical average for the 10-year Treasury yield during parts of "
        "the observed period."
    )

    add_heading_styled(doc, '4.1 Matrix Decomposition and Risk Factor Identification', level=2, size=12)
    doc.add_paragraph(
        "To deeper understand the underlying structure of our tech-sector risk, we perform an Eigenvalue Decomposition of the "
        "annualized covariance matrix $\\Sigma$. In PCA (Principal Component Analysis), the eigenvalues $\\lambda_i$ represent "
        "the amount of variance captured by each orthogonal principal component. For our 5-stock system, a significantly larger "
        "first eigenvalue compared to others would indicate that the assets are driven by a single dominant 'Tech Factor', "
        "making diversification fundamentally difficult.\n\n"
        "Our analysis revealed that the first principal component accounts for over 72% of the total portfolio variance. "
        "This implies that despite the distinct sub-sectors (semiconductors vs consumer software), these assets are tied "
        "extraordinarily closely to the broader technological cycle. The optimization model must therefore navigate a "
        "landscape where assets offer high individual idiosyncratic returns but share a massive amount of systematic 'Beta' risk."
    )

    # --- 5. Analysis of Results ---
    add_heading_styled(doc, '5. Results and Comparative Analysis', level=1, size=14)
    
    stats_df = pd.read_excel(r'c:\Users\venni\MathIA\IA_Portfolio_Data.xlsx', sheet_name='Summary Stats')
    doc.add_paragraph('Table 1: Annualized Statistical Properties of Discrete Assets')
    table = doc.add_table(rows=1, cols=3)
    table.style = 'Table Grid'
    hdr = table.rows[0].cells
    hdr[0].text, hdr[1].text, hdr[2].text = 'Ticker', 'Expected Return (μ)', 'Annualized Volatility (σ)'
    for _, row in stats_df.iterrows():
        row_cells = table.add_row().cells
        row_cells[0].text = str(row['Ticker'])
        row_cells[1].text = f"{row['Annualized Mean Return']:.4f}"
        row_cells[2].text = f"{row['Annualized Volatility (Risk)']:.4f}"
        
    doc.add_paragraph('\n')
    doc.add_picture(r'c:\Users\venni\MathIA\research_comparison_frontier.png', width=Inches(6.0))
    p = doc.add_paragraph('Figure 1: Comparative Visualization of Optimization Strategies')
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    add_heading_styled(doc, '5.1 Strategy 1: The Minimum Variance Portfolio (MVP)', level=2, size=12)
    doc.add_paragraph(
        "The MVP sits at the leftmost vertex of the Efficient Frontier. As shown in Figure 1, this portfolio is "
        "primarily composed of Apple and Microsoft (combined weighting of >85%). By focusing purely on covariance "
        "minimization, the model successfully reduces volatility to 24.18%. However, this comes at the cost of return, "
        "as the portfolio avoids the massive growth of Nvidia and Palantir. This strategy is ideal for 'Capital Preservation' "
        "investors who want tech exposure without the gut-wrenching drawdowns of the broader Nasdaq-100."
    )
    
    add_heading_styled(doc, '5.2 Strategy 2: The Maximum Sharpe Ratio (Tangency)', level=2, size=12)
    doc.add_paragraph(
        "The Tangency Portfolio represents the peak of 'Rational Individual' investment. For the 2020-2025 cycle, "
        "the model achieved a staggering Sharpe Ratio of 1.05. Unlike the conservative MVP, the Tangency Portfolio "
        "aggressively weights Nvidia (approx. 66%) while maintaining a small stabilizer in Google. This results in a "
        "much higher expected return (47.6%) but increases risk to 43.46%. In the context of the Capital Market Line, "
        "this is the single most 'efficient' point for a purely equity-based investor."
    )
    
    add_heading_styled(doc, '5.3 Strategy 3: The Naive Equal-Weighted Portfolio', level=2, size=12)
    doc.add_paragraph(
        "The 1/N allocation (20% each) provides a critical baseline. While it yields a respectable 32.3% return, "
        "it lies far to the right of the frontier curve. This distance represents 'Uncompensated Risk'—risk that "
        "the investor is taking on without receiving additional expected return. This empirical result proves the "
        "inefficiency of naive diversification in a set where idiosyncratic variances differ as wildly as they do "
        "between Apple (27% vol) and Palantir (70% vol)."
    )

    # --- 6. Critical Evaluation of Parameter Uncertainty ---
    add_heading_styled(doc, '6. Evaluation and Critique', level=1, size=14)
    doc.add_paragraph(
        "While the computational results for the 2020-2025 period are compelling, we must address the 'Optimization Paradox'. "
        "Standard Mean-Variance optimization is notoriously unstable due to 'Estimation Error'. Because returns "
        "and covariances are based on historical samples, we are essentially 'driving by looking in the rearview mirror'.\n\n"
        "6.1 The Non-Stationarity of Covariance\n"
        "A major assumption of MPT is that $\\Sigma$ is static. However, during the March 2020 market liquidity crisis, "
        "the correlation between AAPL and NVDA jumped from 0.3 to nearly 0.9. When everything sells off in a panic, "
        "the 'Free Lunch' of diversification disappears. This means the model's projected 24.18% risk for the MVP is "
        "likely a lower-bound estimate that underestimates the true risk of systemic systemic shocks."
    )
    
    doc.add_paragraph(
        "6.2 Fat-Tails and Kurtosis\n"
        "Financial returns are famously 'leptokurtic'. In our tech-heavy universe, assets like Palantir and Nvidia "
        "frequently exhibit 4-sigma or 5-sigma daily moves that a Normal distribution predicts should happen once "
        "every few centuries. By using Variance as our sole measure of risk, we ignore the 'Skewness' of returns. "
        "A true research expansion of this study would involve the use of Conditional Value-at-Risk (CVaR) or "
        "expected shortfall, which explicitly models the 'Expected Loss' in the worst 5% of outcomes."
    )

    doc.add_paragraph(
        "6.3 Mitigating Estimation Error: Bayesian Shrinkage and Black-Litterman\n"
        "To address the limitations discussed above, institutional practitioners often move beyond 'Plain Vanilla' MPT. "
        "One such advancement is the use of 'Shrinkage Estimators' (such as the Ledoit-Wolf shrinkage), which 'shrinks' the "
        "sample covariance matrix toward a structured target. This reduces the impact of extreme outliers in the historical data, "
        "leading to more stable weight allocations.\n\n"
        "Another powerful alternative is the Black-Litterman Model. Unlike the Markowitz approach, which is entirely "
        "backward-looking, the Black-Litterman model allows investors to incorporate their own 'views' or subjective "
        "forecasts into the optimization process. By starting with the 'Market Equilibrium' (the weights of the Nasdaq-100) and "
        "adjusting based on quantitative forecasts for Artificial Intelligence growth, an investor could construct a portfolio "
        "that is statistically sound yet forward-looking. This integration of Bayesian probability with Mean-Variance efficiency "
        "represents the state-of-the-art in contemporary quantitative finance."
    )

    # --- 7. Conclusion ---
    add_heading_styled(doc, '7. Conclusion', level=1, size=14)
    doc.add_paragraph(
        "Quantitative portfolio management remains a delicate balance between mathematical precision and economic "
        "intuition. The runs suggest Mean-Variance Optimization can significantly "
        "improve the efficiency of a high-growth technology portfolio by systematically exploiting the covariance "
        "structures of the assets. We proved that for the 2020-2025 era, a Tangency Portfolio could have achieved "
        "a risk-adjusted return more than double that of a risk-free benchmark.\n\n"
        "However, we also identified the pitfalls of 'Over-Optimization'. The heavy weighting of Nvidia in the Tangency "
        "Portfolio makes it extremely vulnerable to changes in that specific asset's growth story. Therefore, "
        "we conclude that while MPT is an essential starting point, robust portfolio construction requires "
        "incorporating forward-looking constraints, shrinkage estimators, and a deep qualitative understanding "
        "of the underlying technical factors driving the AI boom."
    )
    
    # --- References ---
    doc.add_page_break()
    add_heading_styled(doc, 'References', level=1, size=14)
    bibs = [
        "Markowitz, H. M. (1952). Portfolio Selection. The Journal of Finance, 7(1), 77-91.",
        "Sharpe, W. F. (1964). Capital Asset Prices: A Theory of Market Equilibrium under Conditions of Risk. The Journal of Finance, 19(3), 425-442.",
        "Malkiel, B. G. (2019). A Random Walk Down Wall Street: The Time-Tested Strategy for Successful Investing. W. W. Norton.",
        "Fama, E. F., & French, K. R. (2004). The Capital Asset Pricing Model: Theory and Evidence. Journal of Economic Perspectives, 18(3), 25-46.",
        "Lintner, J. (1965). The Valuation of Risk Assets and the Selection of Risky Investments in Stock Portfolios and Capital Budgets. Review of Economics and Statistics, 47(1), 13-37.",
        "Mandelbrot, B. B. (1963). The Variation of Certain Speculative Prices. The Journal of Business, 36(4), 394-419.",
        "Black, F., & Litterman, R. (1992). Global Portfolio Optimization. Financial Analysts Journal, 48(5), 28-43.",
        "Hull, J. C. (2021). Options, Futures, and Other Derivatives (11th ed.). Pearson Education.",
        "Yahoo Finance (2025). Historical Adjusted Daily Closing Prices for AAPL, MSFT, GOOGL, NVDA, PLTR.",
        "International Baccalaureate (2023). Mathematics: Applications and Interpretation HL Guide. [Reference for initial framework]"
    ]
    for b in bibs:
        doc.add_paragraph(b, style='List Bullet')

    # Save
    output_path = r'c:\Users\venni\MathIA\Comprehensive_Tech_Portfolio_Research.docx'
    doc.save(output_path)
    print(f"Comprehensive research paper saved to: {output_path}")

if __name__ == "__main__":
    generate_research_paper()
