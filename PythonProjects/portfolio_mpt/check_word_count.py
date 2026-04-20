import docx

def check_word_count(file_path):
    doc = docx.Document(file_path)
    total_words = 0
    for p in doc.paragraphs:
        total_words += len(p.text.split())
    
    # Also check tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                total_words += len(cell.text.split())
                
    return total_words

print(f"Word Count: {check_word_count(r'c:\Users\venni\MathIA\Full_Dissertation_Portfolio_Optimization.docx')}")
