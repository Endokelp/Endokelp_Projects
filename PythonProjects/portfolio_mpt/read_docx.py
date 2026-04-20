import docx
import os

file_path = r'c:\Users\venni\MathIA\MathIA_High_Quality_Final.docx'

if not os.path.exists(file_path):
    print(f"File not found: {file_path}")
else:
    doc = docx.Document(file_path)
    print(f"Reading: {file_path}")
    print(f"Number of paragraphs: {len(doc.paragraphs)}")
    
    with open('draft_content.txt', 'w', encoding='utf-8') as f:
        for i, para in enumerate(doc.paragraphs):
            if para.text.strip():
                f.write(f"{i}: {para.text}\n")
    print("Content written to draft_content.txt")
